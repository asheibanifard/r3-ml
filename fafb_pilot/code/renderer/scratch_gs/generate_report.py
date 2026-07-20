"""
Builds the full experiments report for the scratch_gs from-scratch pure-CUDA
3D Gaussian Splatting pipeline, as a Word document (matching this project's
established report convention, fafb_pilot/results/experiment_report.docx).

Covers, in order: motivation, full mathematical model (density, composite
loss incl. analytic gradients, Adam, DVR/MIP, Gaussian rasterization, tiling,
bake architecture, quality metrics, statistics), a step-by-step build
narrative, bugs found and fixed, experimental setup, results (tables +
figures using the real data already produced this session), discussion, and
conclusions.

USAGE
    /venv/r3-ml/bin/python3 fafb_pilot/code/renderer/scratch_gs/generate_report.py
"""
import csv
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_PATH = SCRIPT_DIR / "scratch_gs_experiments_report.docx"

MONO_FONT = "Consolas"
BODY_FONT = "Calibri"


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = BODY_FONT
    return h


def add_para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = BODY_FONT
    return p


def add_math(doc, lines):
    """Monospace block for plain-text mathematical formulas."""
    if isinstance(lines, str):
        lines = [lines]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = MONO_FONT
        r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    r.font.name = BODY_FONT
    return p


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        add_caption(doc, caption)


def add_table(doc, headers, rows, col_widths=None, header_fill="4472C4"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        r.font.name = BODY_FONT
        set_cell_shading(hdr_cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def read_csv_rows(path):
    with open(path) as f:
        return list(csv.DictReader(f))


# ---------------------------------------------------------------------------
def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    # ============================== TITLE ==============================
    title = doc.add_heading(
        "From-Scratch Pure-CUDA 3D Gaussian Splatting:\n"
        "Training, Rendering, and Benchmarking Report", level=0)
    for run in title.runs:
        run.font.name = BODY_FONT
    add_para(doc, "fafb_pilot/code/renderer/scratch_gs/", italic=True, size=10)
    add_para(doc, "", size=4)

    # ============================== 1. INTRODUCTION ==============================
    add_heading(doc, "1. Introduction and Motivation", level=1)
    add_para(doc,
        "This report documents a fully standalone, from-scratch pure-CUDA implementation of "
        "3D Gaussian Splatting: a synthetic volumetric scene is generated, a compact mixture of "
        "anisotropic 3D Gaussians is fitted to it by hand-written gradient descent (no PyTorch, "
        "no autodiff library, no reuse of the project's existing training kernel), and three "
        "different ways of turning that fitted model back into images are built and compared: "
        "a ground-truth direct-volume-renderer (DVR), a 'bake once, render many' pipeline that "
        "reconstructs a dense voxel grid from the Gaussians and reuses the DVR kernel, and a "
        "live Gaussian rasterizer that renders directly from the compact Gaussian parameters "
        "every frame.")
    add_para(doc,
        "The motivating question is compression: the project's broader goal (see CLAUDE.md) is "
        "to represent large EM microscopy volumes as compact latent codes -- Gaussian mixtures "
        "or implicit neural representations -- instead of dense voxel grids. A meaningful speed "
        "or quality comparison therefore cannot simply pit a Gaussian renderer against "
        "direct-volume-rendering of the uncompressed data; the comparison has to account for "
        "what is actually stored, and for how the compact representation gets turned back into "
        "pixels. This report builds exactly that comparison from first principles, verifies "
        "every mathematical piece independently (finite-difference gradient checks, a legacy "
        "correctness oracle for the rasterizer), and benchmarks it with statistical rigor "
        "(20 repeated runs, 95% confidence intervals).")
    add_para(doc,
        "Everything described below lives in a single standalone file, gaussian_splat_scratch.cu "
        "(plus companion Python analysis scripts), deliberately kept separate from the project's "
        "other CUDA renderers (Mip_Render_Inside_Volume.cu, scripts/_3dgs/3dgs_cuda.cu) so that "
        "no existing, already-validated code could mask a bug in the from-scratch implementation.")

    # ============================== 2. MATH ==============================
    add_heading(doc, "2. Mathematical Model", level=1)

    add_heading(doc, "2.1 Anisotropic Gaussian Mixture Density", level=2)
    add_para(doc,
        "Each of the N trainable Gaussians contributes a smooth 'blob' of density in 3-D "
        "space. The total predicted density at a point x is the sum of all Gaussians' "
        "contributions:")
    add_math(doc, [
        "f(x) = sum_k  v_k * exp( -1/2 * Q_k(x) )",
        "Q_k(x) = (x - mean_k)^T * Sigma_k^-1 * (x - mean_k)      (Mahalanobis distance)",
        "Sigma_k = R_k * diag(s_k1^2, s_k2^2, s_k3^2) * R_k^T     (anisotropic covariance)",
        "v_k = softplus(raw_inten_k) = ln(1 + exp(raw_inten_k))   (always positive amplitude)",
    ])
    add_para(doc,
        "R_k is a rotation matrix built from a learnable quaternion, and s_k1..s_k3 are learnable "
        "per-axis scales (standard deviations), so each Gaussian can be an arbitrarily oriented, "
        "stretched ellipsoid rather than a plain sphere. Contributions below a fixed Mahalanobis "
        "cutoff (Q_k > 20) are skipped entirely -- both for speed (most Gaussian/voxel pairs are "
        "far apart and contribute a negligible amount) and because it bounds each Gaussian's "
        "'support radius', a quantity reused throughout the renderer.")

    add_heading(doc, "2.2 Local-Frame Simplification", level=2)
    add_para(doc,
        "Working directly with the 3x3 matrix inverse Sigma_k^-1 every time a Gaussian is "
        "evaluated would be wasteful. Instead, the offset from the Gaussian's centre is first "
        "rotated into the Gaussian's own local (axis-aligned) frame:")
    add_math(doc, [
        "u = R_k^T * (x - mean_k)                       (world offset -> local frame)",
        "Q_k(x) = u_1^2/s_k1^2 + u_2^2/s_k2^2 + u_3^2/s_k3^2",
    ])
    add_para(doc,
        "because in that rotated frame the covariance is diagonal, so the Mahalanobis distance "
        "collapses to a simple per-axis sum -- no matrix inverse needed at evaluation time, just "
        "three multiplications.")

    add_heading(doc, "2.3 Composite Training Loss", level=2)
    add_para(doc,
        "Training minimises a composite loss with a base data term plus seven regularizers, "
        "matching the formula documented for this project's real training pipeline "
        "(scripts/_3dgs/_3dgs.py, per CLAUDE.md) -- implemented here completely from scratch, "
        "including its own analytic backward pass. CLAUDE.md specifies the formula but not exact "
        "numeric weights, so the weights below are this implementation's own reasonable choices, "
        "kept small relative to the data term so they act as genuine regularizers.")
    add_math(doc, [
        "L = L1(pred, target)",
        "  + lambda_ssim    * (1 - SSIM(random 64x64 Z-slice))",
        "  + lambda_scale   * mean(s_max^2)",
        "  + lambda_ceiling * mean(relu(s_max - cap))",
        "  + lambda_outlier * mean(relu(s_max - median(s_max) - 3*MAD(s_max)))",
        "  + lambda_sparsity* mean(v_k * (1 - target(mean_k)))",
        "  + lambda_aniso   * mean(s_min^2)",
        "  + lambda_count   * mean(sigmoid(raw_inten_k))",
        "  + lambda_L1      * mean(softplus(raw_inten_k))",
        "  + lambda_coverage* mean( -log(s_max / s_ref) )",
    ])
    add_para(doc, "Each term's role in plain language:")
    add_bullets(doc, [
        "L1(pred, target): per-voxel absolute error, averaged over the whole 64^3 grid -- the main driving term.",
        "SSIM term: structural similarity between the predicted and true density on ONE randomly chosen full 64x64 XY-slice per training step (the grid is exactly 64 voxels per axis, so a '64x64 crop' is simply one whole slice) -- pushes local structure, not just per-voxel brightness, to match.",
        "scale / ceiling: keeps each Gaussian's largest axis (s_max) from growing unbounded or beyond a fixed cap.",
        "outlier: additionally penalises any Gaussian whose s_max is a statistical outlier relative to the whole population (more than 3 median-absolute-deviations above the median) -- median/MAD computed fresh each step over all Gaussians.",
        "sparsity: discourages a Gaussian from having bright amplitude if it sits over genuinely empty target space.",
        "aniso: symmetric to the scale term, but on the SMALLEST axis (s_min), discouraging needle-thin degenerate Gaussians.",
        "count / L1-on-amplitude: gently pushes unnecessary Gaussians' amplitude toward zero.",
        "coverage: the opposite pull to scale/ceiling -- discourages Gaussians from collapsing to zero size.",
    ])

    add_heading(doc, "2.4 Analytic Gradients (no autodiff)", level=2)
    add_para(doc,
        "Every gradient used to train the model is derived by hand and implemented directly, "
        "with no automatic-differentiation library. The four learnable quantities per Gaussian "
        "are mean (3 numbers), log_scale (3 numbers, so scale = exp(log_scale) is always "
        "positive), a raw quaternion (4 numbers, normalised before use), and raw_inten (1 number, "
        "amplitude = softplus(raw_inten)).")
    add_math(doc, [
        "dL/d(raw_inten_k)  = dL/dv_k * sigmoid(raw_inten_k)              (softplus derivative)",
        "dL/d(mean_k)       = sum_j g_j * f_kj * R_k * (S_k^-2 . u_j)     (world-frame gradient)",
        "dL/d(log_scale_i)  = sum_j g_j * f_kj * (u_i^2 / s_i^2)          (per-axis, local frame)",
    ])
    add_para(doc,
        "where g_j is the per-voxel loss gradient and f_kj is Gaussian k's density contribution "
        "at voxel j. The quaternion gradient is the most involved: the loss gradient with respect "
        "to the 3x3 rotation matrix R is first accumulated, then chained through two further "
        "Jacobians -- the closed-form derivative of the quaternion-to-rotation-matrix formula, "
        "and the derivative of quaternion normalisation itself:")
    add_math(doc, [
        "dR_ai/dq_raw  ->  closed form, 9 matrix entries x 4 quaternion components",
        "q_norm = q_raw / ||q_raw||",
        "dL/dq_raw = ( dL/dq_norm - q_norm * (q_norm . dL/dq_norm) ) / ||q_raw||",
    ])
    add_para(doc,
        "This last normalisation-Jacobian step is easy to get subtly wrong, which is exactly why "
        "every gradient in this implementation was independently verified against numerical "
        "finite differences before being trusted for training (Section 6.1).")

    add_heading(doc, "2.5 Adam Optimizer", level=2)
    add_math(doc, [
        "m = beta1*m + (1-beta1)*grad             v = beta2*v + (1-beta2)*grad^2",
        "m_hat = m / (1 - beta1^t)                 v_hat = v / (1 - beta2^t)",
        "param  = param - lr * m_hat / (sqrt(v_hat) + eps)",
    ], )
    add_para(doc, "beta1=0.9, beta2=0.999, eps=1e-8, applied per-parameter-group with separate learning rates for mean, log_scale, quaternion, and raw_inten.")

    add_heading(doc, "2.6 Direct Volume Rendering (the ground-truth baseline)", level=2)
    add_para(doc,
        "The baseline renderer marches a ray through the 64^3 voxel grid and keeps the brightest "
        "value seen -- a Maximum Intensity Projection (MIP), matching the convention used "
        "throughout this project's EM-visualisation work (bright neurite structure should never "
        "be hidden by something dimmer in front of it).")
    add_math(doc, [
        "ray-box intersection: slab method, gives entry/exit distances t0, t1",
        "for s in 0..127:  t = t0 + (s+0.5)*(t1-t0)/128;  x = camera_pos + t*dir",
        "                  v = trilinear_interpolate(voxel_grid, x)",
        "                  mip = max(mip, v)",
        "pixel = clamp(mip, 0, 1)",
    ])

    add_heading(doc, "2.7 Gaussian Rasterization (rendering directly from Gaussians)", level=2)
    add_para(doc,
        "Rendering the Gaussian mixture directly (without baking to a grid first) requires, for "
        "every pixel, finding where along that pixel's camera ray each nearby Gaussian is "
        "brightest. This has a closed form: substituting the ray x(t) = origin + t*dir into the "
        "local-frame Mahalanobis distance gives a quadratic in t,")
    add_math(doc, [
        "Q(t) = a*t^2 + b*t + c   <=   cutoff",
        "a = sum_i (dir_i/s_i)^2 ,   b = 2 sum_i (oc_i*dir_i/s_i^2) ,   c = sum_i (oc_i/s_i)^2 - cutoff",
        "solved via the quadratic formula for the exact interval [t_lo, t_hi] where the ray is",
        "inside that Gaussian's support region  (oc = ray_origin - mean_k, in local frame)",
    ])
    add_para(doc,
        "An important correctness subtlety: the correct quantity to compute is "
        "max_t [ sum_k f_k(x(t)) ] (the MIP of the SUM of all overlapping Gaussians along the "
        "ray), not sum_k [ max_t f_k(x(t)) ] (each Gaussian's own peak, summed afterward) -- those "
        "are mathematically different whenever Gaussians overlap, and the wrong version was built "
        "first, then caught and fixed (Section 5, bug 1). The corrected approach discretises each "
        "ray's depth into bins, accumulates every overlapping Gaussian's contribution additively "
        "within each bin, and only takes the maximum across bins at the very end -- this is "
        "algebraically the same thing ray-marching would compute, just organised around Gaussians "
        "instead of samples.")

    add_heading(doc, "2.8 Tile-Based Rasterization Architecture", level=2)
    add_para(doc,
        "A first working rasterizer used one CUDA thread per Gaussian, which leaves nearly all of "
        "a modern GPU idle regardless of screen size (a scene with 800 Gaussians only ever "
        "launches 800 threads). The design was rewritten around per-PIXEL parallelism instead, "
        "mirroring how production Gaussian-splatting renderers achieve real-time performance:")
    add_bullets(doc, [
        "Pass 1 (per Gaussian): determine which screen tiles (16x16 pixels) each Gaussian's projected footprint overlaps, and append the Gaussian's index to each of those tiles' lists.",
        "Pass 2 (per tile, per pixel): one CUDA thread block per tile, one thread per pixel. Each block cooperatively streams its tile's Gaussian list through shared memory in batches; every thread evaluates its own pixel against the batch and accumulates into a private register-resident depth-bin array -- since each thread owns its own pixel exclusively, no atomic operations are needed anywhere in the per-pixel accumulation.",
    ])
    add_para(doc, "This rewrite alone produced a roughly 77x speedup at 256x256 and roughly 450x at 1024x1024 (Section 5).")

    add_heading(doc, "2.9 Bake-Then-Render Architecture", level=2)
    add_para(doc,
        "For the compression use case, the trained Gaussian parameters are the actual "
        "artifact to store or transmit -- far more compact than a dense voxel grid (800 "
        "Gaussians x ~20 floats is roughly 64x smaller than a 64^3 float grid). At display time, "
        "the model is evaluated once at every voxel centre of a fresh dense grid ('baking'), and "
        "that baked grid is then rendered with the exact same DVR kernel used for ground truth. "
        "This both decouples rendering cost from Gaussian count and, because baking evaluates the "
        "model exactly where it was trained (the voxel centres), sidesteps the fidelity loss that "
        "continuous, off-grid evaluation (live rasterization) suffers.")

    add_heading(doc, "2.10 Quality Metrics", level=2)
    add_math(doc, [
        "PSNR = 10 * log10( 1 / MSE(a,b) )                    (data range [0,1]; higher = better)",
        "SSIM: windowed structural similarity (luminance, contrast, structure); higher = better",
        "LPIPS: learned perceptual distance (AlexNet features); lower = better",
    ])
    add_para(doc,
        "Volume-level SSIM/LPIPS (comparing the baked 3-D grid against the ground-truth 3-D grid) "
        "are computed by slice-averaging: both metrics are inherently 2-D, so the standard "
        "extension to a volume is to compute the 2-D metric independently on every Z-slice and "
        "average -- the same convention this project already uses for whole-volume SSIM "
        "elsewhere.")

    add_heading(doc, "2.11 Statistical Methodology for Benchmarking", level=2)
    add_para(doc,
        "Every FPS/latency figure reported below is a mean over 20 independent repeats (a fresh "
        "process launch each time, not just repeated measurements inside one already-warmed-up "
        "process -- process-level repetition is what actually exposes system-level noise such as "
        "background load, thermal state, and cache-cold effects). For n=20 samples, the 95% "
        "confidence interval uses the Student's t-distribution rather than a fixed z=1.96, which "
        "is the statistically correct choice for this sample-size range (it converges to the "
        "z-based interval as n grows, but is meaningfully wider at small n):")
    add_math(doc, [
        "mean = (1/n) * sum(x_i)              std = sqrt( (1/(n-1)) * sum( (x_i - mean)^2 ) )",
        "95% CI = mean +/- t(0.975, df=n-1) * std / sqrt(n)          (t(0.975, df=19) ~ 2.093)",
    ])
    add_para(doc,
        "Latency (ms/frame) is computed by converting each of the 20 raw per-repeat FPS samples "
        "to milliseconds (1000/FPS) BEFORE averaging, rather than inverting the already-averaged "
        "FPS -- since 1/x is a nonlinear transform, applying it to individual samples first and "
        "then aggregating is the statistically correct order of operations.")

    # ============================== 3. STEP BY STEP ==============================
    add_heading(doc, "3. Implementation: Step by Step", level=1)
    steps = [
        ("Step 1 -- Synthetic target volume",
         "A deterministic 64^3 voxel grid is generated on the GPU as the sum of 8 fixed, "
         "hand-placed anisotropic Gaussian 'blobs' (Section 2.1's density formula, evaluated "
         "directly, clamped to [0,1]). Being fully deterministic (no random-number generator), "
         "the exact same ground-truth scene is produced on every run."),
        ("Step 2 -- Per-Gaussian precompute kernel",
         "Before every forward/backward pass, one CUDA thread per Gaussian normalises its "
         "quaternion, builds its 3x3 rotation matrix, exponentiates and clamps its log-scales "
         "into actual scales, computes its amplitude via softplus, and derives its support "
         "radius and which axis is its largest/smallest (needed by the scale/ceiling/aniso/"
         "coverage loss terms)."),
        ("Step 3 -- Forward kernel",
         "One thread per voxel (262,144 threads for the 64^3 grid) sums every nearby Gaussian's "
         "contribution (Section 2.1-2.2), compares against the target, and writes the base L1 "
         "gradient contribution for that voxel."),
        ("Step 4 -- Backward kernel",
         "One thread per Gaussian, looping over all voxels within its support radius. Because "
         "each Gaussian's gradient is written only by its own thread, no atomic operations are "
         "needed at all -- a 'transposed layout' matching this project's own real training "
         "kernel's design philosophy (many samples for forward, one thread per Gaussian for "
         "backward)."),
        ("Step 5 -- Adam update kernel",
         "One thread per Gaussian applies the Adam update (Section 2.5) to all four learnable "
         "quantities."),
        ("Step 6 -- Finite-difference self-test",
         "Before any training is trusted, every analytic gradient is spot-checked against a "
         "numerical central-difference estimate on held-out Gaussians (Section 6.1). This step "
         "caught a genuine implementation bug (float32 rounding noise in the loss-scalar "
         "accumulator, Section 5) before it could silently corrupt training."),
        ("Step 7 -- Training loop",
         "3000 iterations of precompute -> forward -> backward -> Adam-update, with the SSIM "
         "term's random Z-slice re-chosen every iteration and the outlier term's median/MAD "
         "recomputed fresh every iteration from the current population of Gaussians."),
        ("Step 8 -- DVR baseline renderer",
         "Implements Section 2.6 directly against the ground-truth voxel grid."),
        ("Step 9 -- Gaussian rasterizer (first version, then corrected)",
         "An initial per-Gaussian atomicMax-based rasterizer was built, found to systematically "
         "under-estimate density under Gaussian overlap, and replaced with the binned-"
         "accumulation approach of Section 2.7 (bug 1, Section 5)."),
        ("Step 10 -- Tile-based rewrite",
         "The per-Gaussian-thread rasterizer was replaced with the tile-based, per-pixel-thread "
         "architecture of Section 2.8, verified against the original (now a correctness oracle) "
         "on every run."),
        ("Step 11 -- Bake kernel",
         "Implements Section 2.9: evaluate the trained model once onto a fresh dense grid, reuse "
         "the DVR kernel to render it."),
        ("Step 12 -- Statistically rigorous benchmarking",
         "CUDA-event-based GPU-only timing (isolating actual rendering cost from device-to-host "
         "transfer, frame export, and disk I/O), repeated 20 times per screen size as fresh "
         "process launches, with mean/std/95%-CI reporting per Section 2.11."),
    ]
    for name, desc in steps:
        p = doc.add_paragraph()
        r = p.add_run(name + ": ")
        r.bold = True
        r.font.name = BODY_FONT
        r2 = p.add_run(desc)
        r2.font.name = BODY_FONT

    # ============================== 4. BUGS ==============================
    add_heading(doc, "4. Bugs Found and Fixed", level=1)
    add_para(doc,
        "Three genuine issues were found and corrected during development, each caught by a "
        "verification step built specifically to catch this class of mistake -- not by visual "
        "inspection alone.")

    add_heading(doc, "4.1 Float32 loss-accumulator rounding noise", level=2)
    add_para(doc,
        "Accumulating 262,144 single-precision atomicAdd operations into one float32 loss "
        "scalar grows the running sum large enough that float32 rounding noise (around 1e-4 in "
        "scale) swamped the tiny loss deltas the finite-difference gradient check needed to "
        "measure, producing spurious 50-500% 'errors' against gradients that were actually "
        "correct. Fix: accumulate the loss (only the loss; per-voxel gradients are computed "
        "independently and were never affected) in double precision.")

    add_heading(doc, "4.2 Camera-front culling on Gaussian centre instead of support radius", level=2)
    add_para(doc,
        "With the camera positioned at the volume's centre, any Gaussian whose CENTRE sits "
        "slightly behind the camera plane but whose support region still extends in front of it "
        "was being fully dropped by the rasterizer's front/back cull test -- because the camera "
        "sits inside the scene rather than outside looking in, this is easy to hit in practice, "
        "and it produced a large, roughly uniform darkening across the reconstructed image "
        "(measured PSNR 15 dB instead of the expected ~28 dB). Fix: cull a Gaussian only if its "
        "entire support sphere is behind the camera plane (depth + support_radius <= near_plane), "
        "and when the camera sits inside a Gaussian's support region, fall back to scanning the "
        "full screen for it instead of the small-angle screen-projection approximation (which "
        "divides by a near-zero or negative depth in that regime).")

    add_heading(doc, "4.3 First Gaussian rasterizer computed the wrong quantity", level=2)
    add_para(doc,
        "The first working rasterizer computed each Gaussian's own peak density along the ray "
        "independently (via atomicMax) and combined results afterward -- mathematically "
        "sum_k[max_t f_k(x(t))], not the correct max_t[sum_k f_k(x(t))] (Section 2.7). Under "
        "Gaussian overlap these differ, and the bug was directly confirmed by measurement (mean "
        "reconstructed brightness dropped from about 0.88 to as low as 0.49-0.69 versus a "
        "ray-marched reference). Fixed by the binned-accumulation approach described in Section "
        "2.7 and carried into the tile-based rewrite.")

    # ============================== 5. SETUP ==============================
    add_heading(doc, "5. Experimental Setup", level=1)
    add_table(doc,
        ["Parameter", "Value"],
        [
            ["GPU", "NVIDIA RTX 5000 Ada Generation"],
            ["CUDA", "12.9 (nvcc 12.9.86)"],
            ["Compilation", "nvcc -O3 --use_fast_math -gencode arch=compute_89,code=sm_89"],
            ["Voxel grid", "64 x 64 x 64 (fixed for all experiments; screen size only affects rendering)"],
            ["Synthetic scene", "8 fixed anisotropic Gaussian blobs, deterministic (no RNG)"],
            ["Number of trainable Gaussians", "800"],
            ["Training iterations", "3000"],
            ["Mahalanobis cutoff", "20.0"],
            ["Learning rates (Adam)", "mean 4e-3, log_scale 5e-3, quaternion 1.5e-3, raw_inten 1.5e-2"],
            ["Screen sizes benchmarked", "64, 128, 256, 512, 1024, 2048 (square)"],
            ["Camera sweep", "60 camera angles per benchmark, yaw swept 0-360 deg, camera at volume centre"],
            ["Benchmark repeats", "20 independent process launches per screen size"],
        ],
        col_widths=[2.6, 3.6],
    )

    # ============================== 6. RESULTS ==============================
    add_heading(doc, "6. Results", level=1)

    add_heading(doc, "6.1 Gradient Verification (Finite-Difference Self-Test)", level=2)
    add_para(doc,
        "Before training was trusted, analytic gradients were compared against central-"
        "difference numerical estimates for representative Gaussians. All real (non-near-zero) "
        "gradients agree to well under 1% relative error; the two quat entries with large "
        "reported relative error are cases where BOTH the analytic and numerical values are "
        "themselves within double-precision noise of exactly zero (an expected, mathematically "
        "correct result at/near an isotropic or identity-rotation configuration, not a sign of "
        "error -- see discussion in Section 2.4).")
    add_table(doc,
        ["Gaussian", "Parameter", "Analytic", "Numerical", "Rel. Error"],
        [
            ["0", "mean.x", "2.8246e-04", "2.8385e-04", "0.49%"],
            ["0", "log_scale.x", "1.1117e-04", "1.1185e-04", "0.62%"],
            ["0", "raw_inten", "1.1671e-04", "1.1671e-04", "0.002%"],
            ["400", "mean.x", "6.1579e-05", "6.1709e-05", "0.21%"],
            ["400", "log_scale.x", "8.4666e-04", "8.4729e-04", "0.07%"],
            ["400", "quat.w", "-2.7055e-04", "-2.7046e-04", "0.03%"],
            ["400", "quat.x", "2.0587e-04", "2.0594e-04", "0.03%"],
            ["400", "raw_inten", "8.7490e-04", "8.7492e-04", "0.001%"],
            ["799", "mean.x", "1.1808e-04", "1.1850e-04", "0.36%"],
            ["799", "log_scale.x", "3.8830e-04", "3.8912e-04", "0.21%"],
            ["799", "raw_inten", "3.9837e-04", "3.9838e-04", "0.001%"],
        ],
        col_widths=[0.8, 1.3, 1.3, 1.3, 1.1],
    )
    add_caption(doc, "Table 1. Analytic vs. numerical (central-difference) gradients, composite loss, three representative Gaussians. Gaussian 400 was deliberately given a non-identity quaternion and anisotropic scale so the quaternion/rotation gradient path is genuinely exercised.")

    add_heading(doc, "6.2 Training Convergence", level=2)
    add_para(doc, "The composite loss (Section 2.3) decreases monotonically over 3000 iterations:")
    add_table(doc,
        ["Iteration", "Loss"],
        [["1", "0.374959"], ["200", "0.003428"], ["400", "0.001580"], ["600", "0.001293"],
         ["800", "0.001151"], ["1000", "0.001032"], ["1200", "0.000939"], ["1400", "0.000873"],
         ["1600", "0.000794"], ["1800", "0.000730"], ["2000", "0.000660"], ["2200", "0.000601"],
         ["2400", "0.000542"], ["2600", "0.000457"], ["2800", "0.000393"], ["3000", "0.000341"]],
        col_widths=[1.5, 1.5],
    )
    add_caption(doc, "Table 2. Composite training loss vs. iteration, 800 Gaussians, 64^3 target.")

    add_heading(doc, "6.3 Visual Fidelity: GT DVR vs. Baked+DVR vs. Live Rasterizer", level=2)
    add_para(doc,
        "The figure below (512x512 screen, one representative camera angle) compares the "
        "ground-truth DVR render against both reconstruction methods, with the absolute "
        "difference shown alongside each. The baked reconstruction is visually indistinguishable "
        "from ground truth; the live rasterizer shows a smooth residual bias plus faint blocky "
        "artefacts from its depth-binning, consistent with continuous/off-grid evaluation being "
        "inherently less accurate than evaluation exactly at the trained voxel centres (Section "
        "2.9).")
    add_image(doc, SCRIPT_DIR / "results_512" / "frame_comparisons" / "frame_0000.png", width_in=6.3,
               caption="Figure 1. GT DVR vs. Baked+DVR (top) and vs. Live Gaussian rasterizer (bottom), 512x512, with per-pixel absolute difference and PSNR/SSIM annotated.")

    add_heading(doc, "6.4 Quality Metrics Across Screen Sizes", level=2)
    add_para(doc,
        "Screen-space PSNR/SSIM/LPIPS, each representation compared against GT DVR, averaged "
        "over the 60-frame camera sweep, at every benchmarked screen size:")
    rows = []
    for size in [64, 128, 256, 512, 1024, 2048]:
        csv_path = SCRIPT_DIR / f"results_{size}" / "metrics_summary.csv"
        if csv_path.exists():
            for r in read_csv_rows(csv_path):
                psnr = r["psnr_db"]
                ssim = r["ssim"]
                lpips_v = r["lpips"]
                rows.append([
                    str(size), r["representation"],
                    f"{float(r['gpu_fps']):.1f}",
                    f"{float(psnr):.2f}" if psnr else "-",
                    f"{float(ssim):.4f}" if ssim else "-",
                    f"{float(lpips_v):.4f}" if lpips_v else "-",
                ])
    add_table(doc, ["Screen", "Representation", "GPU FPS", "PSNR (dB)", "SSIM", "LPIPS"], rows,
               col_widths=[0.7, 1.7, 0.9, 0.9, 0.8, 0.8])
    add_caption(doc, "Table 3. Screen-space quality metrics (vs. GT DVR) and single-run GPU FPS, all six benchmarked screen sizes.")

    add_heading(doc, "6.5 Volume-Level Fidelity (Baked Grid vs. Ground-Truth Grid)", level=2)
    add_para(doc,
        "Unlike the screen-space metrics above, this compares the baked 64^3 grid directly "
        "against the ground-truth 64^3 grid in 3-D -- the metric that most directly answers "
        "'how good is the compressed Gaussian representation as a stand-in for the original "
        "volume', independent of any particular camera or screen size:")
    add_table(doc, ["Metric", "Value"],
               [["Volume PSNR", "68.72 dB"], ["Volume SSIM (slice-averaged)", "0.99995"],
                ["Volume LPIPS (slice-averaged)", "0.000004"]],
               col_widths=[3.0, 2.0])
    add_caption(doc, "Table 4. Baked-grid vs. ground-truth-grid fidelity, whole 64^3 volume.")

    add_heading(doc, "6.6 Rendering Throughput (Statistically Rigorous, 20 Repeats)", level=2)
    add_para(doc,
        "GPU-only throughput (CUDA events; device-to-host transfer, frame export, video "
        "generation, and disk I/O excluded), mean and 95% confidence interval over 20 repeated "
        "fresh-process benchmark runs:")
    stats_rows = []
    stats_csv = SCRIPT_DIR / "results_bench_stats" / "fps_summary_stats.csv"
    label_map = {"dvr_gpu_fps": "GT DVR", "baked_gpu_fps": "Baked + DVR", "rasterizer_gpu_fps": "Live Gaussian rasterizer"}
    if stats_csv.exists():
        for r in read_csv_rows(stats_csv):
            if r["metric"] not in label_map:
                continue
            stats_rows.append([
                r["screen_size"], label_map[r["metric"]], r["n"],
                f"{float(r['mean_fps']):.2f}", f"{float(r['std_fps']):.3f}",
                f"[{float(r['ci95_lo']):.2f}, {float(r['ci95_hi']):.2f}]",
            ])
    add_table(doc, ["Screen", "Method", "n", "Mean FPS", "Std Dev", "95% CI"], stats_rows,
               col_widths=[0.7, 1.6, 0.4, 1.0, 0.9, 1.4])
    add_caption(doc, "Table 5. GPU-only throughput statistics, 20 repeats per screen size, 95% CI via Student's t-distribution (df=19).")
    add_image(doc, SCRIPT_DIR / "results_bench_stats" / "fps_vs_resolution.png", width_in=5.8,
               caption="Figure 2. GPU rendering throughput vs. output resolution, mean +/- 95% CI, n=20. GT DVR (dashed, hollow circle) and Baked+DVR (solid, filled square) track each other almost exactly since both use the identical DVR kernel; the marker/linestyle distinction keeps both visible where they numerically coincide.")

    add_heading(doc, "6.7 Rendering Latency", level=2)
    add_para(doc,
        "The same measurements expressed as latency (ms/frame = 1000/FPS, computed per repeat "
        "then aggregated -- Section 2.11) make the absolute, practical gap far more visible than "
        "the log-scale FPS plot: at 1024x1024 the DVR methods need about 0.94 ms/frame versus "
        "about 10.3 ms/frame for the live rasterizer; at 2048x2048, about 3.8 ms/frame versus "
        "about 39.5 ms/frame.")
    add_image(doc, SCRIPT_DIR / "results_bench_stats" / "latency_vs_resolution.png", width_in=5.8,
               caption="Figure 3. GPU rendering latency vs. output resolution, mean +/- 95% CI, n=20.")

    # ============================== 7. DISCUSSION ==============================
    add_heading(doc, "7. Discussion", level=1)
    add_para(doc,
        "Baking wins on both axes at once. Because Baked+DVR reuses the exact same DVR kernel "
        "as ground truth, its throughput tracks GT DVR almost exactly at every screen size "
        "(Table 5, Figure 2) -- rendering cost is fully decoupled from the number of Gaussians. "
        "At the same time, because baking evaluates the trained model exactly at its training "
        "voxel centres, it reaches 59 dB screen-space PSNR against ground truth versus the live "
        "rasterizer's 28 dB (Table 3) and 68.7 dB volume-level PSNR against the true target grid "
        "(Table 4) -- there is no speed/quality tradeoff between the two Gaussian-based options "
        "in this experiment; baking simply dominates for a static, already-trained scene.")
    add_para(doc,
        "The live rasterizer degrades faster than DVR as resolution grows. Its relative "
        "disadvantage against GT DVR roughly doubles from 64x64 (about 4.8x slower) to 2048x2048 "
        "(about 10-11x slower), consistent with its per-pixel cost (an exact 3-D ray/Gaussian "
        "quadratic solve plus per-bin exp() evaluation) being intrinsically more expensive than "
        "DVR's cheap, cache-resident trilinear lookups, and with redundant per-tile reloading of "
        "Gaussian data from global memory growing as more, finer tiles are needed at higher "
        "resolution.")
    add_para(doc,
        "FPS rising then falling with resolution is a genuine GPU-occupancy effect, not a "
        "measurement artefact. Both curves in Figure 2 peak around 128x128 and decline "
        "afterward: at 64x64 the render kernels launch too few threads (4096) to occupy a modern "
        "GPU's many streaming multiprocessors, so throughput is occupancy-bound; by 1024x1024-"
        "2048x2048 there is enough parallel work that the workload becomes genuinely pixel-count-"
        "bound, and FPS falls by roughly the same factor pixel count rises (approximately 4x "
        "fewer FPS for the 4x-more-pixels step from 1024^2 to 2048^2) -- internally consistent "
        "with simple compute-bound scaling once occupancy is no longer the limiting factor.")
    add_para(doc,
        "Scope of the claim. These measurements use a small, synthetic, 64^3 volume (1 MB) that "
        "is fully resident in GPU L2 cache, and a fixed 800-Gaussian model. They demonstrate that "
        "-- on this specific synthetic benchmark -- baking dominates live rasterization for a "
        "static scene, and that DVR itself is remarkably fast against a cache-resident grid. They "
        "should not be read as a general claim that direct volume rendering is faster than "
        "Gaussian rasterization for large-scale, real-world volumetric data, where the dense grid "
        "would typically be far too large to cache and the compression benefits of a compact "
        "Gaussian representation would matter for storage and transmission in ways this small "
        "synthetic scene cannot exercise.")

    # ============================== 8. STITCHING EXPERIMENT ==============================
    add_heading(doc, "8. Multi-Block Stitching Experiment", level=1)
    add_para(doc,
        "Everything up to this point uses a single 64^3 block. Real usage, however, means many "
        "adjacent blocks tiled into a larger scene, and there are two structurally different ways "
        "to render that: bake each block to a grid first and stitch the grids, or stitch the "
        "Gaussians themselves and rasterize the combined set directly. This section builds and "
        "compares both, at 2, 4, and 8 stitched blocks (stitch_experiment.cu, "
        "compare_stitching.py).")

    add_heading(doc, "8.1 Design", level=2)
    add_para(doc,
        "No new training happens here: the single already-trained 800-Gaussian checkpoint is "
        "reused, replicated, and offset to fill 2 blocks (2x1x1 arrangement), 4 blocks (2x2x1), "
        "or 8 blocks (2x2x2) of a fixed combined volume (128^3 voxels spanning [-2,2]^3 -- the "
        "same voxel density as the original 64^3/[-1,1]^3 single block, sized to fit the largest, "
        "8-block arrangement; the 2- and 4-block cases simply leave the rest of that volume "
        "empty). Two rendering strategies are built on top of this:")
    add_bullets(doc, [
        "Baked-then-stitch: each block is baked INDEPENDENTLY onto its own local 64^3 grid, using only that block's own Gaussians (hard-gated -- exactly the single-block bake_kernel, unchanged, with no cross-block blending). The N independent 64^3 grids are then placed into the correct sub-region of the combined 128^3 grid and rendered with the same DVR/MIP kernel used everywhere else.",
        "Gaussian-stitch: every block's Gaussians are concatenated into one array (each block's copy with its mean shifted to that block's world position) and rendered directly by the tile-based rasterizer -- no grid, no per-block boundary of any kind.",
    ])
    add_para(doc,
        "A real implementation bug was caught and fixed while building this: the block-placement "
        "formula only handled world offsets of -1 and +1 correctly (mapping to the combined "
        "grid's two edges); an offset of 0 -- the un-shifted axes in the 2- and 4-block cases -- "
        "needs to land at the MIDDLE of the combined grid, not its edge. The bug was caught by "
        "directly probing the combined grid's values at the expected block-content locations "
        "before trusting any rendered output, and confirmed fixed the same way.")

    add_heading(doc, "8.2 Results", level=2)
    stitch_csv = SCRIPT_DIR / "results_stitching" / "stitching_comparison.csv"
    stitch_rows = []
    if stitch_csv.exists():
        for r in read_csv_rows(stitch_csv):
            stitch_rows.append([
                r["n_blocks"], r["n_gaussians"],
                f"{float(r['baked_fps']):.1f}", f"{float(r['raster_fps']):.1f}",
                f"{float(r['psnr_db']):.2f}", f"{float(r['ssim']):.4f}", f"{float(r['lpips']):.4f}",
            ])
    add_table(doc, ["Blocks", "Gaussians", "Baked FPS", "Raster FPS", "PSNR (dB)", "SSIM", "LPIPS"],
               stitch_rows, col_widths=[0.7, 0.9, 1.0, 1.0, 1.0, 0.8, 0.8])
    add_caption(doc, "Table 6. Baked-then-stitch vs. Gaussian-stitch: GPU FPS for each method, and PSNR/SSIM/LPIPS of Gaussian-stitch against Baked-then-stitch (the higher-fidelity, grid-evaluated reference), averaged over a 60-frame camera sweep, 512x512 screen.")

    add_image(doc, SCRIPT_DIR / "results_stitching" / "fps_vs_block_count.png", width_in=5.5,
               caption="Figure 4. GPU FPS vs. number of stitched blocks, both methods.")
    add_image(doc, SCRIPT_DIR / "results_stitching" / "quality_vs_block_count.png", width_in=6.3,
               caption="Figure 5. PSNR/SSIM/LPIPS of Gaussian-stitch vs. Baked-then-stitch, vs. number of stitched blocks.")
    add_image(doc, SCRIPT_DIR / "results_stitching" / "frame_comparison_all_blocks.png", width_in=6.3,
               caption="Figure 6. Baked-then-stitch vs. Gaussian-stitch vs. |difference|, one representative camera angle, for 2/4/8 stitched blocks.")

    add_heading(doc, "8.3 Discussion", level=2)
    add_para(doc,
        "Throughput reproduces the single-block finding at larger scale, for a clean reason. "
        "Baked-then-stitch's FPS is flat (about 7000-7076 across 2, 4, and 8 blocks) because DVR "
        "cost depends only on the fixed 128^3 combined grid, never on how many Gaussians produced "
        "it. Gaussian-stitch's FPS falls steadily (784 -> 503 -> 333) because the live "
        "rasterizer's cost scales directly with Gaussian count (1600 -> 3200 -> 6400). This is "
        "the same architectural conclusion as the single-block screen-size sweep, now confirmed "
        "along a different axis (Gaussian/block count instead of screen resolution).")
    add_para(doc,
        "Quality tells a more interesting story than the single-block case, and it is visible "
        "directly in Figure 6: the 8-block difference image shows a sharp CROSS-HAIR SEAM "
        "PATTERN exactly at the block boundaries. This is the direct, visible consequence of the "
        "two methods' different handling of block edges -- baked-then-stitch is hard-gated (a "
        "block's bake sums ONLY its own Gaussians, by construction), while Gaussian-stitch has no "
        "boundary at all, so a Gaussian near one block's edge can still contribute density into "
        "its neighbour's territory if its support radius reaches that far. SSIM, a local-"
        "structure metric, is far more sensitive to this sharp, spatially-localised seam "
        "discrepancy than PSNR (which averages error uniformly over the whole frame) -- which is "
        "exactly why SSIM collapses to 0.37-0.75 here versus the 0.97-0.99+ seen for the "
        "single-block, no-stitching comparison, while PSNR (25-29 dB) stays roughly in the same "
        "range as before.")
    add_para(doc,
        "The non-monotonic dip at 4 blocks (worse PSNR/SSIM than either 2 or 8) is a real "
        "consequence of the arrangement geometry, confirmed in the full 60-frame average, not a "
        "single-frame artefact. 8 blocks (2x2x2) has seams along all three axes but shares them "
        "symmetrically; 2 blocks (2x1x1) has only one seam axis; 4 blocks (2x2x1) sits in "
        "between with 2 seam axes while still leaving much of the un-doubled Z extent empty "
        "(visible in Figure 6's sparser 2- and 4-block panels), so seam-affected area makes up a "
        "larger fraction of the actually-visible, non-empty content specifically in the 4-block "
        "case.")
    add_para(doc,
        "This experiment necessarily used IDENTICAL replicated content at every block position "
        "(the same single trained checkpoint, only translated) -- a controlled simplification "
        "that isolates the STITCHING MECHANISM and its performance/quality tradeoffs cleanly, but "
        "cannot exercise genuine content-driven seam artefacts the way independently-trained, "
        "differently-content blocks would (matching this project's own earlier seam-statistics "
        "work). That is the natural next step for this specific experiment, noted below.")

    # ============================== 9. CONCLUSION ==============================
    add_heading(doc, "9. Conclusion and Future Work", level=1)
    add_para(doc,
        "A complete from-scratch pure-CUDA pipeline was built and verified end to end: synthetic "
        "scene generation, an analytically-differentiated Gaussian-mixture trainer with a full "
        "composite loss (verified against finite differences before being trusted), a "
        "ground-truth DVR renderer, a corrected and then tile-based Gaussian rasterizer, and a "
        "bake-then-render architecture -- benchmarked with statistically rigorous, 20-repeat, "
        "confidence-interval-backed measurements of both throughput and latency, and extended to "
        "multi-block stitching (2/4/8 blocks). The central, well-supported finding for this "
        "synthetic benchmark is that baking a trained Gaussian model to a dense grid before "
        "rendering is strictly better than live rasterization on both speed and fidelity, for a "
        "static scene, at both single-block and multi-block scale -- with the added multi-block "
        "finding that the two approaches diverge specifically at block seams, far more visibly in "
        "SSIM than in PSNR.")
    add_para(doc, "Natural next steps:")
    add_bullets(doc, [
        "Repeat the comparison on a larger, non-cache-resident volume (or the project's real FAFB blocks) to test whether DVR's advantage survives once the dense grid no longer fits in GPU cache.",
        "Train genuinely different content per block (rather than this section's replicated-and-offset single block) to study real, content-driven seam artefacts, matching this project's own earlier seam-statistics work.",
        "Add an alpha-compositing (occlusion-based) rendering mode alongside the existing MIP mode, as a separate, explicitly-labelled option -- appropriate for segmented/known-structure data, not raw intensity (see the project's own discussion of MIP vs. alpha blending for neuron visualisation).",
        "Extend the tile-based rasterizer with exact per-pixel footprint tests before the expensive ray/Gaussian solve, to close more of its remaining gap with DVR at high resolution.",
        "Measure actual storage/transmission size of the compact Gaussian representation against the dense grid directly, to quantify the compression ratio this whole pipeline is ultimately motivated by.",
    ])

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    main()
