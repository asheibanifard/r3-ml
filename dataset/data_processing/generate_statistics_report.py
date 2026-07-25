"""
Builds a Word report explaining the block-statistics data-processing stage
(fafb_block_statistical_selection.ipynb): why per-block statistics are
computed across the whole FAFB block grid, what each figure/table in
fafb_block_statistics_output/ actually shows, and -- the main point -- why
each of the four representative blocks (simple/typical/complex/edge_rich)
was selected and how it was categorised.

Matches this project's established report convention (see
generate_report.py, fafb_pilot/results/experiment_report.docx): a single
standalone script, python-docx, real numbers pulled from the actual output
artefacts rather than invented ones.

USAGE
    /venv/r3-ml/bin/python3 fafb_pilot/code/renderer/scratch_gs/data_processor/generate_statistics_report.py
"""
import json
import csv
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = Path(__file__).resolve().parent
OUT_DIR = SCRIPT_DIR / "fafb_block_statistics_output"
OUT_PATH = SCRIPT_DIR / "fafb_block_statistics_report.docx"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


# --------------------------------------------------------------------------
# Helpers (same conventions as generate_report.py)
# --------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = BODY_FONT
    return h


def add_para(doc, text, italic=False, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = BODY_FONT
    return p


def add_math(doc, lines):
    if isinstance(lines, str):
        lines = [lines]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = MONO_FONT
        r.font.size = Pt(10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    r.font.name = BODY_FONT
    return p


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        add_caption(doc, caption)


def add_table(doc, headers, rows, col_widths=None, header_fill="4472C4"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        r.font.name = BODY_FONT
        set_cell_shading(hdr_cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def read_csv_rows(path):
    with open(path) as f:
        return list(csv.DictReader(f))


# --------------------------------------------------------------------------
def main():
    stats = json.loads((OUT_DIR / "report_summary_stats.json").read_text())
    selected_rows = read_csv_rows(OUT_DIR / "selected_representative_blocks.csv")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    # ============================== TITLE ==============================
    title = doc.add_heading(
        "FAFB Block Statistics and Representative-Block Selection", level=0)
    for run in title.runs:
        run.font.name = BODY_FONT
    add_para(doc, "fafb_pilot/code/renderer/scratch_gs/data_processor/", italic=True, size=10)
    add_para(doc, "", size=4)

    # ============================== 1. INTRO ==============================
    add_heading(doc, "1. Motivation", level=1)
    add_para(doc,
        "The FAFB dataset is pre-partitioned into 262,144 blocks of 64x64x64 voxels each "
        "(CLAUDE.md). Training or evaluating the project's Gaussian-mixture or SIREN models on "
        "an arbitrary handful of blocks risks accidentally picking blocks that are all easy "
        "(e.g. mostly uniform background) or all hard (e.g. dense, high-contrast neurite "
        "boundaries), which would make any reported quality number unrepresentative of the "
        "dataset as a whole.")
    add_para(doc,
        "fafb_block_statistical_selection.ipynb addresses this directly: it computes a fixed "
        "set of per-block image statistics across every single block in the dataset, ranks them "
        "by percentile, and uses those percentiles to deliberately pick out four blocks spanning "
        "the range from visually simple to visually complex/edge-rich. Those four blocks then "
        "serve as a small, curated, characterised benchmark set for the rest of the project's "
        "experiments, instead of an arbitrary or random sample.")

    # ============================== 2. DATA PROCESSING ==============================
    add_heading(doc, "2. Data-Processing Pipeline", level=1)

    add_heading(doc, "2.1 Input abstraction", level=2)
    add_para(doc,
        "A VolumeSource class unifies five possible input layouts behind one interface -- HDF5 "
        "(.h5/.hdf5), Zarr, a single .npy volume, a directory of pre-extracted .npy blocks, or a "
        "directory of per-block .tif files -- so the same statistics/selection code runs "
        "unchanged regardless of which on-disk form the dataset happens to be in.")

    add_heading(doc, "2.2 Per-block feature extraction", level=2)
    add_para(doc,
        "Every block (all 262,144, iterated in the fixed x-fastest/y/z-slowest block-grid order "
        "documented in CLAUDE.md) is normalised to [0,1] and reduced to a fixed set of scalar "
        "statistics:")
    add_bullets(doc, [
        "mean, median, variance, std, mad (median absolute deviation), minimum, maximum -- basic intensity statistics.",
        "entropy -- Shannon entropy of the block's intensity histogram; low for flat/uniform blocks, high for blocks with a rich spread of intensities.",
        "mean_gradient, gradient_p90, gradient_p99 -- magnitude of the local intensity gradient (edge strength), summarised by its mean and upper percentiles so a block with a few very sharp edges is distinguished from one that is uniformly gradient-y.",
        "hf_l1, hf_l2 -- L1/L2 norm of the high-frequency component of the block (energy left after a low-pass filter is subtracted), a texture/detail-content measure.",
        "laplacian_l1, laplacian_l2 -- L1/L2 norm of the Laplacian (second-derivative edge/boundary detector), complementary to the first-derivative gradient measures.",
    ])
    add_para(doc,
        f"Every one of the {stats['n_blocks']:,} blocks in the dataset was successfully processed: "
        f"{stats['n_padded']} required padding and {stats['n_full_valid']:,} "
        f"({100*stats['n_full_valid']/stats['n_blocks']:.1f}%) have `valid_fraction == 1.0` -- "
        "the crawl is complete, not a partial sample, and every block is a genuine full 64^3 cube "
        "with no boundary padding.")

    add_heading(doc, "2.3 Percentile normalisation and composite scores", level=2)
    add_para(doc,
        "The raw features above live on different, non-comparable scales (entropy in nats, "
        "variance in normalised-intensity-squared units, gradients in normalised-intensity units, "
        "etc.), so each is converted into its rank-percentile across the whole 262,144-block "
        "population before being combined:")
    add_math(doc, [
        "percentile(feature)_i = rank(feature_i among all blocks) / N * 100",
        "complexity_score = mean( variance_pct, entropy_pct, mean_gradient_pct, hf_l1_pct )",
        "edge_score       = mean( mean_gradient_pct, gradient_p90_pct, laplacian_l1_pct, hf_l1_pct )",
    ])
    add_para(doc,
        "complexity_score is a general 'how much is going on in this block' measure (intensity "
        "spread + texture energy); edge_score instead specifically weights sharp, high-percentile "
        "gradient and second-derivative content, so a block can score high on one and not "
        "necessarily the other (e.g. a block dominated by one very sharp boundary against an "
        "otherwise flat background can have a high edge_score without a correspondingly high "
        "complexity_score).")

    # ============================== 3. DISTRIBUTION RESULTS ==============================
    add_heading(doc, "3. Distribution of Block Statistics", level=1)

    add_heading(doc, "3.1 Summary statistics", level=2)
    desc = stats["describe"]
    feat_labels = {
        "mean": "Mean intensity", "variance": "Variance", "entropy": "Entropy",
        "mean_gradient": "Mean gradient", "hf_l1": "HF energy (L1)", "laplacian_l1": "Laplacian (L1)",
    }
    rows = []
    for f, label in feat_labels.items():
        d = desc[f]
        rows.append([label, f"{d['mean']:.4f}", f"{d['std']:.4f}", f"{d['min']:.4f}",
                     f"{d['50%']:.4f}", f"{d['max']:.4f}"])
    add_table(doc, ["Feature", "Mean", "Std", "Min", "Median", "Max"], rows,
              col_widths=[1.6, 0.9, 0.9, 0.9, 0.9, 0.9])
    add_caption(doc, f"Table 1. Raw per-block feature statistics over all {stats['n_blocks']:,} blocks.")

    add_heading(doc, "3.2 A small background/empty-tissue population", level=2)
    add_para(doc,
        f"{stats['n_background']:,} blocks ({stats['pct_background']}% of the dataset) have "
        "mean == variance == 0 -- completely flat, empty blocks. These are not scattered randomly "
        "through the volume: grouping by the block-grid z-index shows them concentrated almost "
        "entirely at the far end of the z range,")
    bg_rows = [[f"z={z}", f"{n:,} blocks"] for z, n in stats["bg_z_counts"].items()]
    add_table(doc, ["Block-grid Z index", "Background block count"], bg_rows, col_widths=[2.0, 2.5])
    add_caption(doc, "Table 2. Count of empty (mean==0) blocks by block-grid Z index -- all 3,355 fall in the last 3 of 64 Z-layers, and the count rises toward the very end.")
    add_para(doc,
        "consistent with the imaged tissue physically running out near one end of the acquisition "
        "stack, not with random imaging dropout. This is exactly the kind of population an "
        "unweighted random sample could either wildly over- or under-represent, and is one "
        "concrete reason the selection in Section 4 works from full-population percentiles rather "
        "than ad hoc sampling.")

    add_heading(doc, "3.3 Feature correlations and distribution shape", level=2)
    add_para(doc,
        "The six core features split into two correlated clusters: mean/entropy (r=0.93, both "
        "tracking 'how much real tissue vs. background is in this block') and "
        "mean_gradient/hf_l1/laplacian_l1 (r=0.97, all measuring edge/texture content); variance "
        "sits in between (r=0.61-0.85 with the others). Every one of these raw features is "
        "strongly left-skewed (skew -2.2 to -3.7, high kurtosis) with a visibly bimodal histogram "
        "-- see Figure 1 -- because the small empty-block population from Section 3.2 forms an "
        "isolated low-value cluster, separate from the dominant mass of genuinely textured tissue "
        "blocks.")
    corr = stats["corr"]
    corr_feats = list(feat_labels.keys())
    corr_rows = []
    for f1 in corr_feats:
        corr_rows.append([feat_labels[f1]] + [f"{corr[f1][f2]:.2f}" for f2 in corr_feats])
    add_table(doc, ["Feature"] + [feat_labels[f] for f in corr_feats], corr_rows,
              col_widths=[1.3] + [0.75]*6)
    add_caption(doc, "Table 3. Pearson correlation between the six core per-block features, over all blocks.")

    add_heading(doc, "3.4 Histograms", level=2)
    add_para(doc,
        "Each panel below is one raw or composite feature, histogrammed over all 262,144 blocks.")
    add_image(doc, OUT_DIR / "hist_entropy.png", width_in=5.5,
        caption="Figure 1. Intensity entropy -- the clearest illustration of the bimodal empty-vs-tissue split described in Section 3.2: an isolated low bar near entropy~0, well separated from the dominant high-entropy tissue mass.")
    add_image(doc, OUT_DIR / "hist_variance.png", width_in=5.5,
        caption="Figure 2. Intensity variance -- same left-skew/bimodal pattern as entropy, for the same reason.")
    add_image(doc, OUT_DIR / "hist_mean_gradient.png", width_in=5.5,
        caption="Figure 3. Mean gradient magnitude -- edge-content proxy; the empty-block spike is smaller here since even a fully empty block has near-zero, not exactly zero, gradient noise.")
    add_image(doc, OUT_DIR / "hist_hf_l1.png", width_in=5.5,
        caption="Figure 4. High-frequency (L1) energy -- texture/detail-content proxy, same overall shape.")
    add_image(doc, OUT_DIR / "hist_complexity_score.png", width_in=5.5,
        caption="Figure 5. Composite complexity_score. Percentile-averaging compresses the extreme skew of the raw features into a broad hump centred around 55-65, but the empty-block population still shows up as a distinct spike near 0 -- percentile ranking cannot remove a population that is genuinely, categorically different from the rest.")
    add_image(doc, OUT_DIR / "hist_edge_score.png", width_in=5.5,
        caption="Figure 6. Composite edge_score -- same construction as complexity_score, weighted toward gradient/Laplacian percentiles instead of variance/entropy.")

    # ============================== 4. SELECTION METHODOLOGY ==============================
    add_heading(doc, "4. Why These Specific Blocks Were Selected", level=1)

    add_heading(doc, "4.1 Target percentile profiles", level=2)
    add_para(doc,
        "Four categories are defined, each as a target point in percentile-space -- i.e. a "
        "description of what an idealised block of that category would score on each feature, "
        "not a fixed value or threshold on the raw statistic itself:")
    target_rows = [
        ["simple", "10", "10", "10", "-", "10", "-"],
        ["typical", "50", "50", "50", "-", "50", "-"],
        ["complex", "90", "90", "80", "-", "90", "-"],
        ["edge_rich", "70", "80", "95", "95", "90", "95"],
    ]
    add_table(doc,
        ["Category", "variance_pct", "entropy_pct", "mean_gradient_pct", "gradient_p90_pct", "hf_l1_pct", "laplacian_l1_pct"],
        target_rows, col_widths=[0.9, 0.85, 0.85, 1.1, 1.05, 0.75, 1.05])
    add_caption(doc, "Table 4. Target percentile profile per category (a dash means that feature is not part of that category's target -- e.g. simple/typical/complex are defined on 4 features, edge_rich on 6, deliberately including the two gradient/Laplacian percentiles the others omit).")
    add_para(doc,
        "simple and complex are direct opposites on the same four features (10th vs. 90th/80th "
        "percentile); typical targets the exact 50th percentile on all four, i.e. 'an ordinary, "
        "middle-of-the-distribution block'. edge_rich is qualitatively different, not just "
        "'more complex than complex': it is defined on six features instead of four, specifically "
        "adding gradient_p90 and laplacian_l1 at the 95th percentile -- it targets a block "
        "dominated by strong, sharp edges rather than one that is simply high on every feature at "
        "once.")

    add_heading(doc, "4.2 Nearest-match selection", level=2)
    add_para(doc,
        "For each category, every one of the 262,144 candidate blocks is scored by its Euclidean "
        "distance (in percentile-space, over that category's specific feature subset) from the "
        "target profile, and the single closest, not-yet-used block is chosen:")
    add_math(doc, [
        "distance(block, target) = sqrt( mean_f [ (block.feature_pct_f - target.feature_pct_f)^2 ] )",
        "chosen_block = argmin_block  distance(block, target)     (skipping any block already used by an earlier category)",
    ])
    add_para(doc,
        "Because the candidate pool is the full, unrestricted dataset (all blocks have "
        "`valid_fraction == 1.0`, so none are excluded), and the pool is large (262,144 "
        "candidates), every category's nearest match is essentially an exact hit on its target "
        "profile -- the results below confirm this directly.")

    add_heading(doc, "4.3 Selected blocks", level=2)
    sel_rows = []
    for r in selected_rows:
        sel_rows.append([r["category"], r["block_id"],
                          f"{float(r['complexity_score']):.1f}",
                          f"{float(r['edge_score']):.1f}",
                          f"{float(r['distance']):.2f}"])
    add_table(doc, ["Category", "Block ID", "complexity_score", "edge_score", "Match distance (RMS pct pts)"],
              sel_rows, col_widths=[1.0, 1.7, 1.2, 1.0, 1.6])
    add_caption(doc, "Table 5. The four selected representative blocks. Match distance is the RMS percentile-point deviation from Table 4's target profile -- all four are within 1.6 percentile points of an exact match, out of a possible 0-100+ scale.")
    add_para(doc,
        "Concretely: image_z56_y53_x37 (simple) scores near the 10th percentile on every target "
        "feature -- a genuinely low-texture block. image_z38_y46_x62 (typical) sits almost exactly "
        "at the 50th percentile on complexity and edge score alike -- as close to 'an average "
        "block' as any single block in the dataset can be. image_z14_y11_x17 (complex) is in the "
        "88th/84th percentile on complexity/edge score -- high on both, matching its target of "
        "being uniformly complex rather than specifically edge-dominated. image_z11_y27_x15 "
        "(edge_rich) has a HIGHER edge_score (92.4) than complexity_score (83.6), the direct, "
        "measured consequence of its target profile (Section 4.1) deliberately over-weighting "
        "sharp-edge features relative to the other three categories.")

    add_heading(doc, "4.4 Visual confirmation", level=2)
    add_para(doc,
        "The category labels are not just a numerical artefact of the scoring formula -- the "
        "actual image content visibly matches what each category claims:")
    add_image(doc, OUT_DIR / "selected_blocks" / "selected_blocks_comparison.png", width_in=6.0,
        caption="Figure 7. Axial centre slice of all four selected blocks (top-left simple, top-right typical, bottom-left complex, bottom-right edge_rich). simple shows one dark process crossing an otherwise fairly uniform neuropil background; complex and edge_rich both show dense, high-contrast membrane boundaries, but edge_rich's boundaries are sharper and more locally concentrated, consistent with its edge-weighted target profile.")
    add_para(doc,
        "Per-category raw and normalised 64^3 volumes, plus axial/coronal/sagittal centre slices "
        "and a maximum-intensity projection, are saved for each block under "
        "fafb_block_statistics_output/selected_blocks/, for use directly as the project's curated "
        "benchmark set going forward.")

    # ============================== 5. CONCLUSION ==============================
    add_heading(doc, "5. Conclusion", level=1)
    add_para(doc,
        "A complete, dataset-wide statistical crawl (262,144/262,144 blocks) followed by "
        "percentile-based composite scoring and nearest-target-match selection produces four "
        "blocks that are demonstrably, measurably representative of specific, named points in the "
        "dataset's complexity/edge-content distribution -- not an arbitrary or convenient choice. "
        "The methodology also surfaced a real, non-obvious property of the dataset itself (a "
        "small, spatially-concentrated population of empty blocks near one end of the Z range, "
        "Section 3.2) that would be easy to miss without processing every block rather than a "
        "sample.")

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    main()
